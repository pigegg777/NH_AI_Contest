from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_NAME = "Malgun Gothic"
TITLE_COLOR = RGBColor(22, 54, 90)
HEADING_COLOR = RGBColor(46, 116, 181)
HEADING_DARK = RGBColor(31, 77, 120)
TEXT_COLOR = RGBColor(34, 34, 34)
MUTED_COLOR = RGBColor(90, 90, 90)
LIGHT_FILL = "EEF4FB"
MILD_FILL = "F7F9FC"
BORDER_COLOR = "D8DEE8"

OUTPUT_NAME = "(구현Track) 에이전트 부문 아이디어 기획서_제출용_구현반영.docx"


def set_run_font(
    run,
    *,
    size=None,
    bold=None,
    color=None,
    italic=None,
):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER_COLOR, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        edge_el = tc_borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:color"), color)


def set_cell_text(
    cell,
    text,
    *,
    bold=False,
    size=10.5,
    color=TEXT_COLOR,
    align=WD_ALIGN_PARAGRAPH.LEFT,
):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=0, after=2, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)


def set_table_widths(table, widths):
    table.autofit = False
    for column, width in zip(table.columns, widths):
        for cell in column.cells:
            cell.width = Inches(width)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED_COLOR)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def apply_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_COLOR

    for style_name, size, color, bold in [
        ("Heading 1", 16, HEADING_COLOR, True),
        ("Heading 2", 12.5, HEADING_COLOR, True),
        ("Heading 3", 11.5, HEADING_DARK, True),
        ("List Bullet", 10.5, TEXT_COLOR, False),
        ("List Number", 10.5, TEXT_COLOR, False),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_run = footer_p.add_run(
        "AI 경진대회 제출용 기획서 | react-app 구현 반영 초안"
    )
    set_run_font(footer_run, size=9, color=MUTED_COLOR)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_page_number(footer.add_paragraph())


def add_title_paragraph(doc, text, *, size=22, after=6, color=TITLE_COLOR):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=0, after=after, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def add_body_paragraph(doc, text, *, before=0, after=6, bold=False, color=TEXT_COLOR):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=after, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_spacing(
        paragraph,
        before=16 if level == 1 else 10,
        after=6 if level == 1 else 4,
        line=1.0,
    )
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=16 if level == 1 else 12.5 if level == 2 else 11.5,
        bold=True,
        color=HEADING_COLOR if level < 3 else HEADING_DARK,
    )
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, before=0, after=4, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=TEXT_COLOR)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(paragraph, before=0, after=4, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=TEXT_COLOR)
    return paragraph


def add_note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_FILL)
    set_cell_border(cell, color="BFD0E5")
    set_cell_text(cell, text, size=10.5)
    return table


def add_summary_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [1.55, 4.95])
    for label, value in rows:
        cells = table.add_row().cells
        shade_cell(cells[0], MILD_FILL)
        set_cell_text(cells[0], label, bold=True, size=10.2, color=HEADING_DARK)
        set_cell_text(cells[1], value, size=10.2)
    return table


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, widths)

    header_cells = table.rows[0].cells
    for index, text in enumerate(headers):
        shade_cell(header_cells[index], LIGHT_FILL)
        set_cell_text(
            header_cells[index],
            text,
            bold=True,
            size=10.2,
            color=HEADING_DARK,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=10.0)

    return table


def add_cover_page(doc):
    add_title_paragraph(doc, "구현 Track 경진대회 기획서", size=24, after=3)
    add_title_paragraph(doc, "Agent 부문", size=16, after=12, color=HEADING_COLOR)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(intro, before=0, after=18, line=1.15)
    run = intro.add_run(
        "실무 효율성과 현장 활용성을 중심으로, 현재 react-app 구현 흐름을 반영해 제출용 문안으로 다듬은 제안서"
    )
    set_run_font(run, size=10.5, color=MUTED_COLOR)

    add_heading(doc, "1. 개요", level=1)
    add_summary_table(
        doc,
        [
            (
                "에이전트명",
                "AI Agent 기반 사무소 농자재 데이터 정제 및 안내 페이지 자동 생성 시스템",
            ),
            (
                "한 줄 소개",
                "사무소 엑셀 가격표를 업로드하면 데이터 이상을 먼저 검토하고, 사무소별 공개 상품 안내 페이지와 QR을 자동 생성·업데이트하는 현장형 에이전트",
            ),
            ("해결 분야", "Work (반복·비효율 업무를 AI로 지원)"),
            (
                "주요 사용자",
                "생산경제·자재 담당 직원, 사무소 관리자, 현장 고객 응대 인력",
            ),
            (
                "구현 상태",
                "로그인, 엑셀 정제, AI 검토 추천, 스토어프런트 빌더, 공개 페이지, QR 출력까지 시연 가능한 프로토타입",
            ),
        ],
    )

    add_body_paragraph(doc, "")
    participant_table = add_table(
        doc,
        ["성명", "소속", "사번 / 역할"],
        [
            ("안호건", "발안농협", "132400608 / 기획·개발"),
            ("", "", ""),
        ],
        [1.5, 2.1, 2.9],
    )
    participant_table.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body_paragraph(doc, "")
    add_note_box(
        doc,
        "핵심 메시지: 이 제안은 단순한 아이디어 소개가 아니라, react-app에 구현된 데이터 정제·사무소별 저장·AI 보조 디자인·공개 페이지·QR 배포 흐름을 실제 업무 문제 해결 프레임으로 재정리한 제출용 기획서이다.",
    )
    add_body_paragraph(
        doc,
        "작성 메모: 제출 전 참가자 정보, 정량 KPI 수치, 사무소명 등은 실제 제출 기준에 맞춰 최종 보정하면 된다.",
        before=8,
        after=0,
        color=MUTED_COLOR,
    )


def add_problem_page(doc):
    doc.add_page_break()
    add_heading(doc, "2. 문제인식 및 필요성", level=1)

    add_heading(doc, "현재 업무의 주요 문제", level=2)
    for item in [
        "가격 변동이나 판매 정책 변경이 생기면 기존 엑셀, 가격표, 안내물을 다시 손봐야 해 현장 반영 속도가 느리다.",
        "기존 가격표는 단가 중심 정보에 치우쳐 있어 성분, 사용량, 특장점, 주의사항 같은 설명 정보가 부족하다.",
        "계통 등록 전체 상품이 한 번에 노출되면 실제 사무소에서 취급하는 품목만 선별해 안내하기 어렵다.",
        "엑셀 데이터에 중복 등록, 가격 역전, 누락 정보가 섞여 있어 저장 전 검토 시간이 많이 든다.",
        "종이 안내물은 한 번 출력되면 최신 가격이나 구성 변경을 즉시 반영하기 어렵다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "왜 지금 필요한가", level=2)
    add_body_paragraph(
        doc,
        "현재 react-app은 로그인, 엑셀 업로드, 데이터 정규화, 검토 추천, 사무소별 저장, 공개 페이지 구성, QR 출력까지 하나의 흐름으로 연결되어 있다. 즉, 이번 제안은 막연한 구상보다 이미 작동하는 프로토타입을 현장 적용 관점에서 구체화하는 작업에 가깝다.",
    )
    add_body_paragraph(
        doc,
        "농자재 업무는 가격 변동 빈도가 높고, 현장에서는 빠른 설명과 신뢰 가능한 최신 정보가 중요하다. 따라서 데이터 정제와 안내 페이지 생성을 분리하지 않고 한 흐름으로 묶는 것이 실질적 효과를 크게 만든다.",
    )

    add_heading(doc, "에이전트 도입 시 기대 변화", level=2)
    for item in [
        "수정 업무를 엑셀 재편집 중심에서 업로드-검토-저장-재배포 중심으로 단순화한다.",
        "사무소별 실제 취급 품목만 골라 안내해 고객 응대 정확도와 설명 속도를 높인다.",
        "AI를 완전 자동 결정이 아니라 검토 우선순위 제안과 디자인 초안 생성에 사용해 실무 적용성을 높인다.",
        "공개 페이지와 QR을 통해 종이 안내물의 한계를 줄이고 최신 정보 접근성을 높인다.",
    ]:
        add_bullet(doc, item)


def add_feature_page(doc):
    doc.add_page_break()
    add_heading(doc, "3. 핵심 기능 및 서비스 흐름", level=1)

    add_heading(doc, "핵심 기능", level=2)
    add_table(
        doc,
        ["기능", "현재 구현 내용", "AI 역할", "실무 효과"],
        [
            (
                "로그인 및 권한 분리",
                "사번 기반 로그인, office_code 기준 데이터 식별, 인증 사용자 세션 유지",
                "직접 생성보다 보안 전제 역할",
                "같은 시스템을 여러 사무소가 써도 자기 데이터만 관리 가능",
            ),
            (
                "엑셀 업로드·정규화",
                "시트 구조를 분석해 헤더·데이터 범위를 찾고 상품 행을 정규화",
                "현재는 규칙 기반 파이프라인 중심",
                "수작업 전처리 없이 원본 가격표를 바로 검토 가능한 구조로 전환",
            ),
            (
                "검토 추천",
                "경고 규칙과 OpenAI 보조 추천을 함께 제공",
                "중복 의심, 가격 이상치, 확인 필요 항목 제안",
                "전체 행을 다시 읽지 않고 우선 점검 포인트부터 확인 가능",
            ),
            (
                "사무소별 데이터 저장",
                "카테고리별 상품 데이터를 office_product_datas에 저장·재호출",
                "AI 없이 검증 가능한 상태 저장",
                "사무소별 운영 데이터 재사용과 업데이트 이력 관리가 쉬워짐",
            ),
            (
                "스토어프런트 빌더",
                "노출 필드 선택, 숨김 상품 관리, 페이지/카드 스타일 구성, 실시간 미리보기 제공",
                "자연어 요청을 구조화된 스타일 의도로 변환",
                "디자인 작업 시간을 줄이고 페이지 품질 편차를 완화",
            ),
            (
                "공개 페이지 및 QR 배포",
                "office_code 기반 공개 URL 렌더링, 링크 복사, PNG 다운로드, 인쇄용 QR 제공",
                "직접 생성보다는 최종 결과물 연결",
                "현장에서 최신 상품 정보를 모바일로 즉시 안내 가능",
            ),
        ],
        [1.15, 2.15, 1.45, 2.05],
    )

    add_heading(doc, "서비스 처리 흐름", level=2)
    for item in [
        "1단계: 사용자가 사무소 계정으로 로그인하고 엑셀 가격표를 업로드한다.",
        "2단계: 시스템이 시트 구조를 분석해 상품 행을 추출하고 경고 항목을 만든다.",
        "3단계: OpenAI 보조 추천이 가능한 환경에서는 추가 검토 포인트를 제안한다.",
        "4단계: 담당자가 카테고리별 데이터를 저장하고 노출 필드와 숨김 상품을 조정한다.",
        "5단계: 페이지 분위기와 카드 구성을 자연어로 수정하며 모바일 미리보기로 확인한다.",
        "6단계: 최종 결과를 공개 페이지와 QR로 연결해 현장 배포에 활용한다.",
    ]:
        add_number(doc, item)

    add_note_box(
        doc,
        "차별점: 단순히 답변을 생성하는 챗봇이 아니라, 엑셀 데이터 정제와 현장용 결과물 생성을 하나의 업무 워크플로우로 연결하는 상태 기반 에이전트라는 점이 핵심이다.",
    )


def add_scenario_page(doc):
    doc.add_page_break()
    add_heading(doc, "4. 사용자 예상 시나리오", level=1)

    scenarios = [
        (
            "시나리오 1. 가격 변동 발생 시 최신 가격을 빠르게 반영",
            "담당 직원이 새 엑셀 파일을 업로드하면 변경된 가격을 반영한 데이터와 안내 페이지를 다시 저장할 수 있다. 이후 공개 페이지는 최신 정보 기준으로 갱신되므로, 종이 가격표를 반복 재출력하지 않고도 QR 중심으로 최신 가격을 안내할 수 있다.",
        ),
        (
            "시나리오 2. 사무소에서 실제 취급하는 품목만 선별해 안내",
            "계통 등록 전체 상품 대신 실제 사무소에서 판매하는 카테고리와 상품만 선택해 고객용 안내 페이지를 구성할 수 있다. 이를 통해 고객은 필요한 상품군만 빠르게 확인하고, 직원은 설명 범위를 줄여 응대 효율을 높일 수 있다.",
        ),
        (
            "시나리오 3. 엑셀 업로드 직후 오류 검토 포인트를 먼저 확인",
            "시스템은 가격 역전, 중복 의심, 핵심 정보 충돌 가능성이 있는 항목을 추천 형태로 제시한다. 담당자는 모든 행을 처음부터 다시 읽지 않고도 우선 점검할 부분부터 확인할 수 있어 현장 반영 전 오류 가능성을 낮출 수 있다.",
        ),
        (
            "시나리오 4. 현장에서 모바일 안내 페이지와 QR로 설명 보조",
            "고객 응대 시 직원은 종이 안내물뿐 아니라 사무소별 모바일 안내 페이지를 함께 보여줄 수 있다. 고객이 QR로 직접 페이지를 열어보면 가격 외에도 상품 특징과 카테고리 구성을 이해할 수 있어 설명 품질이 높아진다.",
        ),
    ]

    for title, body in scenarios:
        add_heading(doc, title, level=2)
        add_body_paragraph(doc, body)


def add_ai_page(doc):
    doc.add_page_break()
    add_heading(doc, "5. AI 활용 계획", level=1)

    add_heading(doc, "AI 활용 영역", level=2)
    add_table(
        doc,
        ["에이전트 역할", "입력", "출력", "현재 상태"],
        [
            (
                "검토 추천 에이전트",
                "정규화된 상품 행, officeCode, 카테고리 정보",
                "중복 의심, 가격 이상치, 확인 필요 추천 목록",
                "구현 완료, OpenAI 연동 시 보조 추천 강화",
            ),
            (
                "페이지 스타일 에이전트",
                "페이지 분위기 요청, 현재 페이지 스타일, 대화 이력",
                "구조화된 JSON 의도와 설명 문구",
                "구현 완료",
            ),
            (
                "카드 디자인 에이전트",
                "노출 필드, 카테고리 샘플, 현재 카드 스타일, 대화 이력",
                "구조화된 JSON 의도와 카드 레이아웃 수정안",
                "구현 완료",
            ),
        ],
        [1.35, 2.2, 1.95, 1.0],
    )

    add_heading(doc, "운영 방식", level=2)
    for item in [
        "OpenAI Responses API의 구조화된 JSON 응답을 받아 프론트엔드 컴파일러가 실제 설정으로 반영한다.",
        "기본 모델은 gpt-4.1-mini 기준으로 분리된 프롬프트를 사용하며, 페이지 스타일·카드 스타일·검토 추천을 목적별로 따로 처리한다.",
        "AI 결과는 곧바로 저장하지 않고, 담당자가 미리보기와 설명을 확인한 뒤 반영 여부를 선택한다.",
        "OpenAI가 없는 환경에서도 규칙 기반 검토 흐름은 독립적으로 동작하도록 설계한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "안전장치", level=2)
    for item in [
        "AI 요청은 officeCode, 현재 스타일, 노출 필드, 상품 정보 등 업무용 비민감 데이터 중심으로 제한한다.",
        "AI 요청 API는 인증된 사용자 세션과 officeCode 소유권 검증을 통과해야만 실행된다.",
        "프롬프트 길이, 요청 본문 크기, 대화 이력 길이를 제한하고, 추천 결과는 자동 수정이 아닌 확인 필요 항목으로만 사용한다.",
    ]:
        add_bullet(doc, item)

    add_note_box(
        doc,
        "왜 AI Agent인가: 이 시스템은 단순 문구 생성기가 아니라, 현재 데이터 상태와 사용자 요청을 함께 읽고, 역할별 초안을 만들고, 사용자 승인 이후에만 결과를 반영하는 상태 기반 업무형 에이전트로 동작한다.",
    )


def add_implementation_page(doc):
    doc.add_page_break()
    add_heading(doc, "6. 구동 방식 및 현재 구현 현황", level=1)

    add_heading(doc, "시스템 구성", level=2)
    add_table(
        doc,
        ["구성 요소", "기술/구현", "설명"],
        [
            (
                "프론트엔드",
                "React 19 + Vite + CSS Modules",
                "로그인, 대시보드, 엑셀 검토 화면, 스토어프런트 빌더, 공개 페이지까지 단일 앱에서 구성",
            ),
            (
                "데이터/인증",
                "Supabase Auth + Database",
                "login_users, office_product_datas, office_page_config, office_page_category_configs 기반 분리 저장",
            ),
            (
                "AI 요청 API",
                "Cloudflare Pages Functions + OpenAI Responses API",
                "페이지 스타일·카드 스타일·검토 추천 요청을 구조화된 JSON으로 중계하고 검증",
            ),
            (
                "공개 배포",
                "공개 URL + qrcode",
                "office_code 기반 공개 링크, QR 다운로드, 링크 복사, 인쇄용 출력 제공",
            ),
            (
                "품질 관리",
                "Vitest 단위 테스트",
                "핵심 모델, AI 응답 계약, 공개 페이지 동작, QR 서비스 등에 대한 회귀 검증 기반 보유",
            ),
        ],
        [1.2, 1.85, 3.45],
    )

    add_heading(doc, "현재 구현 기준 주요 화면", level=2)
    for item in [
        "로그인/회원가입: 사번 기반 계정 생성, 사무소 코드 연동, 세션 유지",
        "상품 데이터 편집 화면: 엑셀 업로드, 경고 확인, 저장, 등록 데이터 재호출",
        "스토어프런트 빌더: 카테고리 선택, 노출 필드 선택, 페이지 스타일, 카드 디자인, 모바일 미리보기",
        "공개 스토어프런트: 사무소별 모바일 상품 안내 페이지 렌더링",
        "QR 출력 기능: 공개 페이지 주소를 PNG/인쇄용 자산으로 생성",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "현재 구현 범위와 차기 확장 범위", level=2)
    add_body_paragraph(doc, "현재 구현 범위", bold=True, after=4)
    for item in [
        "회원가입·로그인 및 사무소별 사용자 식별",
        "엑셀 시트 구조 분석, 상품 행 정규화, 경고 생성",
        "규칙 기반 검토 추천과 OpenAI 보조 추천",
        "사무소별 상품 데이터 저장/불러오기",
        "스토어프런트 빌더와 공개 페이지 렌더링",
        "공개 페이지용 QR 생성 및 출력",
    ]:
        add_bullet(doc, item)

    add_body_paragraph(doc, "차기 확장 범위", bold=True, before=8, after=4)
    for item in [
        "고객용/관리자용 페이지 권한 및 템플릿 고도화",
        "QR 대량 배포와 인쇄 운영 자동화",
        "조회 통계, 추천 반영률, 운영 리포트 분석",
        "비료·농약 외 품목군 확장",
    ]:
        add_bullet(doc, item)


def add_effect_page(doc):
    doc.add_page_break()
    add_heading(doc, "7. 기대 효과 및 KPI", level=1)

    add_heading(doc, "정량·정성 기대 효과", level=2)
    add_table(
        doc,
        ["항목", "기존 방식", "도입 후 목표"],
        [
            (
                "가격표/안내물 수정 시간",
                "엑셀 수정 후 재편집·재출력 중심",
                "업로드-검토-저장 중심으로 전환해 수정 시간 50% 이상 단축 목표",
            ),
            (
                "현장 정보 최신성",
                "출력물 기준으로 최신성 유지가 어려움",
                "공개 페이지와 QR을 통해 최신 정보 중심 안내",
            ),
            (
                "설명 정확도",
                "담당자 개인 숙련도에 따라 편차 발생",
                "사무소별 표준 페이지와 선택 필드 중심 설명",
            ),
            (
                "데이터 검토 효율",
                "전체 행을 수작업으로 다시 확인",
                "AI/규칙 추천으로 우선 검토 포인트부터 점검",
            ),
            (
                "확장성",
                "사무소별 별도 문서 재작업 필요",
                "같은 구조를 여러 사무소에 재사용 가능",
            ),
        ],
        [1.35, 2.35, 2.8],
    )

    add_heading(doc, "운영 KPI 예시", level=2)
    for item in [
        "엑셀 업로드 후 안내 페이지 반영까지 걸리는 평균 시간",
        "AI 추천이 생성한 검토 포인트 중 실제 수정·확인으로 이어진 비율",
        "사무소별 등록 카테고리 수와 공개 페이지 운영 횟수",
        "가격 변동 이후 최신 페이지 갱신 완료까지의 리드타임",
        "QR 페이지 조회 수 및 현장 활용 빈도",
    ]:
        add_bullet(doc, item)

    add_note_box(
        doc,
        "심사 포인트 제안: 본 프로젝트의 핵심 가치는 AI가 멋진 문구를 쓰는 것보다, 실제 반복 업무를 줄이고 최신 정보 전달 속도를 높이는 데 있다. 발표와 문서 모두 이 점을 전면에 두는 것이 좋다.",
    )


def add_expansion_page(doc):
    doc.add_page_break()
    add_heading(doc, "8. 확산 가능성 및 고도화 로드맵", level=1)

    add_heading(doc, "확산 가능성", level=2)
    for item in [
        "같은 형식의 가격표를 사용하는 다른 사무소도 office_code 기준으로 분리 운영할 수 있어 공용 플랫폼화가 쉽다.",
        "사무소별로 노출 카테고리와 페이지 구성을 다르게 관리할 수 있어 공통 플랫폼과 현장 맞춤형 운영을 동시에 만족시킨다.",
        "현재는 비료·농약 중심 흐름이지만, 향후 종자·자재·생활물자 등 다른 품목군으로 확장할 수 있다.",
        "가격표뿐 아니라 행사 안내, 계절 상품 추천, 신규 상품 소개 페이지 등 다른 고객 접점 업무에도 응용 가능하다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "고도화 로드맵", level=2)
    add_table(
        doc,
        ["단계", "목표", "주요 내용"],
        [
            ("1단계", "현재 흐름 안정화", "엑셀 검토 정확도 개선, 사무소별 운영 가이드 정리"),
            ("2단계", "안내 정보 보강", "성분·사용량·주의사항·이미지 링크 연계 고도화"),
            ("3단계", "품목 확장", "종자·자재 등 타 카테고리 데이터 구조 추가"),
            ("4단계", "운영 분석", "조회 통계, 추천 반영률, 자주 수정되는 항목 분석"),
        ],
        [0.9, 1.5, 4.1],
    )


def add_security_page(doc):
    doc.add_page_break()
    add_heading(doc, "9. 운영·보안 준수 및 제출 전 확인", level=1)

    add_heading(doc, "운영 및 보안 원칙", level=2)
    for item in [
        "사무소 단위 권한 분리를 위해 로그인 사용자와 office_code의 소유 관계를 검증한다.",
        "외부 AI 서비스에는 고객 개인정보가 아닌 상품·스타일·검토용 비민감 데이터만 입력하는 것을 원칙으로 한다.",
        "공개 페이지에는 공개 범위로 승인된 정보만 노출하고, 내부 편집 데이터는 인증 영역에서만 다룬다.",
        "실운영 이전에는 AI 요청 로그 범위, 데이터 반출 정책, 공개 페이지 노출 범위를 별도로 점검한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "제출 전 최종 수정 포인트", level=2)
    for item in [
        "참가자 구성, 역할, 제출일을 실제 제출본 기준으로 최종 수정",
        "예상 절감 시간과 KPI 수치를 내부 검토 후 확정",
        "사무소명, 시나리오, 문체를 발표 자료와 일관되게 보정",
        "개인정보 및 보안 준수 여부를 최종 점검",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "개인정보 및 보안 준수 확인", level=2)
    add_note_box(
        doc,
        "본 팀은 과제 수행 과정에서 회사 기밀, 고객 개인정보 등 민감 자료를 외부 AI 서비스에 입력하거나 유출하지 않을 것을 확인합니다. 또한 제출 전 보안 기준과 공개 범위를 다시 점검하겠습니다.",
    )

    add_body_paragraph(doc, "")
    sign_table = doc.add_table(rows=4, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    sign_table.autofit = False
    set_table_widths(sign_table, [1.5, 3.0])
    sign_rows = [
        ("제출일", "2026년     월     일"),
        ("팀장 성명", "(서명)"),
        ("팀원 성명", "(서명)"),
        ("팀원 성명", "(서명)"),
    ]
    for row, (label, value) in zip(sign_table.rows, sign_rows):
        shade_cell(row.cells[0], MILD_FILL)
        set_cell_text(row.cells[0], label, bold=True, color=HEADING_DARK)
        set_cell_text(row.cells[1], value)


def build_document():
    doc = Document()
    apply_document_styles(doc)
    add_cover_page(doc)
    add_problem_page(doc)
    add_feature_page(doc)
    add_scenario_page(doc)
    add_ai_page(doc)
    add_implementation_page(doc)
    add_effect_page(doc)
    add_expansion_page(doc)
    add_security_page(doc)
    return doc


def main():
    root_dir = Path(__file__).resolve().parents[1]
    output_path = root_dir / OUTPUT_NAME
    document = build_document()
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
